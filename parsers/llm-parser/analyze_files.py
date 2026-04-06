"""
Анализ файлов расписаний для понимания структуры
"""
import os
from docx import Document
import win32com.client
import pythoncom

def extract_text_from_docx(file_path):
    """Извлекает текст из .docx файла"""
    try:
        doc = Document(file_path)
        
        print(f"\n{'='*80}")
        print(f"📄 Файл: {os.path.basename(file_path)}")
        print(f"{'='*80}")
        
        # Параграфы
        print(f"\n📝 ПАРАГРАФЫ:")
        for i, para in enumerate(doc.paragraphs[:10]):  # Первые 10
            if para.text.strip():
                print(f"  [{i}] {para.text.strip()}")
        
        # Таблицы
        print(f"\n📊 ТАБЛИЦЫ: {len(doc.tables)}")
        for table_idx, table in enumerate(doc.tables[:2]):  # Первые 2 таблицы
            print(f"\n  Таблица {table_idx + 1}: {len(table.rows)} строк x {len(table.columns)} колонок")
            print(f"  Первые 5 строк:")
            for row_idx, row in enumerate(table.rows[:5]):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"    [{row_idx}] {' | '.join(cells)}")
        
        return True
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        return False

def extract_text_from_doc(file_path):
    """Извлекает текст из .doc файла через Word COM"""
    try:
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(os.path.abspath(file_path))
        
        print(f"\n{'='*80}")
        print(f"📄 Файл: {os.path.basename(file_path)}")
        print(f"{'='*80}")
        
        # Параграфы
        print(f"\n📝 ПАРАГРАФЫ:")
        for i in range(min(10, doc.Paragraphs.Count)):
            para_text = doc.Paragraphs(i+1).Range.Text.strip()
            if para_text:
                print(f"  [{i}] {para_text}")
        
        # Таблицы
        print(f"\n📊 ТАБЛИЦЫ: {doc.Tables.Count}")
        for table_idx in range(min(2, doc.Tables.Count)):
            table = doc.Tables(table_idx + 1)
            print(f"\n  Таблица {table_idx + 1}: {table.Rows.Count} строк x {table.Columns.Count} колонок")
            print(f"  Первые 5 строк:")
            for row_idx in range(min(5, table.Rows.Count)):
                row = table.Rows(row_idx + 1)
                cells = []
                for cell_idx in range(row.Cells.Count):
                    cells.append(row.Cells(cell_idx + 1).Range.Text.strip())
                print(f"    [{row_idx}] {' | '.join(cells)}")
        
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        
        return True
    except Exception as e:
        print(f"❌ Ошибка: {e}")
        try:
            word.Quit()
            pythoncom.CoUninitialize()
        except:
            pass
        return False

if __name__ == "__main__":
    schedules_dir = "../../schedules"
    
    # Анализируем несколько .docx файлов
    print("\n" + "="*80)
    print("АНАЛИЗ .DOCX ФАЙЛОВ")
    print("="*80)
    
    docx_files = [
        "1AP01_03.docx",
        "1GP01.docx",
        "1KP01.docx"
    ]
    
    for filename in docx_files:
        file_path = os.path.join(schedules_dir, filename)
        if os.path.exists(file_path):
            extract_text_from_docx(file_path)
    
    # Анализируем несколько .doc файлов
    print("\n" + "="*80)
    print("АНАЛИЗ .DOC ФАЙЛОВ")
    print("="*80)
    
    doc_files = [
        "1AD01z.doc",
        "1DP01.doc",
        "1EN01z.doc"
    ]
    
    for filename in doc_files:
        file_path = os.path.join(schedules_dir, filename)
        if os.path.exists(file_path):
            extract_text_from_doc(file_path)
