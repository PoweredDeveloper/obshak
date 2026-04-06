"""Просмотр содержимого файла расписания"""
from docx import Document
import sys

file_path = sys.argv[1] if len(sys.argv) > 1 else "../../schedules/1AP01_03.docx"

doc = Document(file_path)

print(f"\n{'='*80}")
print(f"ФАЙЛ: {file_path}")
print(f"{'='*80}\n")

print("ПАРАГРАФЫ:")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

print(f"\n{'='*80}")
print(f"ТАБЛИЦЫ: {len(doc.tables)}")
print(f"{'='*80}\n")

for i, table in enumerate(doc.tables):
    print(f"\nТАБЛИЦА {i+1} ({len(table.rows)} строк x {len(table.columns)} колонок):")
    print("-" * 80)
    for row_idx, row in enumerate(table.rows[:15]):  # Первые 15 строк
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"{row_idx:2d} | {' | '.join(cells)}")
