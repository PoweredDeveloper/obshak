#!/usr/bin/env python3
"""
Парсер расписания КГАСУ из Word файлов
"""
import re
from pathlib import Path
from docx import Document
from datetime import datetime

def parse_time(time_str):
    """Парсит время из строки '8:00:00 - 9:30:00' или '8.00 - 9.30'"""
    if not time_str or '-' not in time_str:
        return None, None
    
    parts = time_str.split('-')
    if len(parts) != 2:
        return None, None
    
    start = parts[0].strip()
    end = parts[1].strip()
    
    # Заменяем точки на двоеточия и добавляем секунды если нужно
    start = start.replace('.', ':')
    end = end.replace('.', ':')
    
    # Добавляем :00 для секунд если их нет
    if start.count(':') == 1:
        start += ':00'
    if end.count(':') == 1:
        end += ':00'
    
    return start, end

def parse_lesson_cell(cell_text):
    """
    Парсит ячейку с занятием
    Формат: 'Математика (Лекции) проф. Тимербаев Н.Ф. 201'
    Может содержать даты: 'до 22.03.' или 'с 06.04.26.'
    """
    if not cell_text or cell_text.strip() == '':
        return None
    
    text = cell_text.strip()
    
    # Извлекаем предмет и тип занятия
    match = re.match(r'(.+?)\s*\((.+?)\)', text)
    if not match:
        return {
            'subject': text,
            'type': None,
            'teacher': None,
            'room': None,
            'start_date': None,
            'end_date': None
        }
    
    subject = match.group(1).strip()
    lesson_type = match.group(2).strip()
    
    # Остальной текст после скобок
    rest = text[match.end():].strip()
    
    # Извлекаем дату окончания если есть (до XX.XX.XX или до XX. XX.)
    end_date = None
    end_date_match = re.search(r'до\s+(\d{2})\s*\.\s*(\d{2})\s*\.?\s*(\d{2})?', rest)
    if end_date_match:
        day = end_date_match.group(1)
        month = end_date_match.group(2)
        year = end_date_match.group(3) if end_date_match.group(3) else '26'
        end_date = f'20{year}-{month}-{day}'
        # Убираем дату из текста
        rest = rest[:end_date_match.start()].strip() + ' ' + rest[end_date_match.end():].strip()
        rest = rest.strip()
    
    # Извлекаем дату начала если есть (С XX.XX.XX. или С XX.XX. или С XX.XX по)
    start_date = None
    # Паттерны для дат начала
    start_patterns = [
        (r'[Сс]\s+(\d{2})\.(\d{2})\.?\s*(\d{2})\.?\s*', lambda m: f'20{m.group(3)}-{m.group(2)}-{m.group(1)}'),
        (r'[Сс]\s+(\d{2})\.(\d{2})\.\s+', lambda m: f'2026-{m.group(2)}-{m.group(1)}'),
        (r'[Сс]\s+(\d{2})\.(\d{2})\s+по\s+', lambda m: f'2026-{m.group(2)}-{m.group(1)}'),
    ]
    
    for pattern, date_func in start_patterns:
        start_match = re.search(pattern, subject)
        if start_match:
            start_date = date_func(start_match)
            # Убираем дату из названия предмета
            subject = re.sub(pattern, '', subject).strip()
            break
    
    # Ищем преподавателя и аудиторию
    teacher = None
    room = None
    
    # Специальные места (спортзалы, лаборатории и т.д.) - обычно в начале или в кавычках
    special_room_match = re.search(r'(СК\s*[«"]?[^»"]+[»"]?|[«"][^»"]+[»"]|Тезуче|НОЦ\s*[«"]?[^»"]+[»"]?)', rest)
    if special_room_match:
        room = special_room_match.group(1).strip()
        # Убираем кавычки если есть
        room = room.replace('«', '').replace('»', '').replace('"', '').strip()
        rest = rest[:special_room_match.start()].strip() + ' ' + rest[special_room_match.end():].strip()
        rest = rest.strip()
    
    # Аудитория обычно в конце (цифры или буквы+цифры, может быть с корпусом типа 2-401)
    if not room:  # Ищем обычную аудиторию только если не нашли специальное место
        # Ищем аудиторию в конце (после удаления дат)
        room_match = re.search(r'(\d+-\d+[А-Яа-я]?|\d+[А-Яа-я]?)\s*$', rest)
        if room_match:
            room = room_match.group(1).strip()
            # Убираем аудиторию
            rest = rest[:room_match.start()].strip()
    
    # Остальное - преподаватель
    if rest:
        teacher = rest
    
    return {
        'subject': subject,
        'type': lesson_type,
        'teacher': teacher,
        'room': room,
        'start_date': start_date,
        'end_date': end_date
    }

def parse_schedule_docx(file_path):
    """Парсит расписание из Word файла"""
    doc = Document(file_path)
    
    if len(doc.tables) == 0:
        print("❌ Таблицы не найдены в документе")
        return []
    
    # Ищем таблицу с расписанием (должна иметь много строк и колонку "Дата")
    table = None
    for t in doc.tables:
        if len(t.rows) > 10:  # Расписание должно иметь много строк
            # Проверяем первую ячейку
            first_cell = t.rows[0].cells[0].text.strip()
            if 'дата' in first_cell.lower():
                table = t
                break
    
    if not table:
        # Если не нашли по критериям, берем самую большую таблицу
        table = max(doc.tables, key=lambda t: len(t.rows))
    
    if len(table.rows) < 2:
        print("❌ Таблица слишком маленькая")
        return []
    
    # Первая строка - заголовки с названиями групп
    header_row = table.rows[0]
    subgroup_row = table.rows[1]
    
    # Собираем информацию о группах и их колонках
    groups_dict = {}
    
    for i, cell in enumerate(header_row.cells[4:], start=4):  # Пропускаем первые 4 колонки
        group_name = cell.text.strip()
        if group_name and re.match(r'\d{1,2}[А-ЯЁ]{2}\d{2}', group_name):
            if group_name not in groups_dict:
                groups_dict[group_name] = []
            
            # Определяем подгруппу из второй строки
            subgroup_text = subgroup_row.cells[i].text.strip().lower()
            subgroup = 0
            if 'подгруппа 1' in subgroup_text:
                subgroup = 1
            elif 'подгруппа 2' in subgroup_text:
                subgroup = 2
            
            # Проверяем не добавили ли мы уже эту подгруппу (для объединенных ячеек)
            column_info = {
                'column_index': i,
                'subgroup': subgroup
            }
            
            # Добавляем только если такой комбинации еще нет
            if not any(c['subgroup'] == subgroup and c['column_index'] != i for c in groups_dict[group_name]):
                groups_dict[group_name].append(column_info)
    
    print(f"📋 Найдено групп: {len(groups_dict)}")
    for group_name, columns in groups_dict.items():
        subgroups = [f"подгр.{col['subgroup']}" for col in columns]
        print(f"  - {group_name}: {', '.join(subgroups)}")
    
    # Парсим расписание
    lessons = []
    day_of_week_map = {
        'понедельник': 1,
        'вторник': 2,
        'среда': 3,
        'четверг': 4,
        'пятница': 5,
        'суббота': 6,
        'воскресенье': 7
    }
    
    current_day = None
    
    for row_idx, row in enumerate(table.rows[2:], start=2):  # Пропускаем заголовки
        cells = row.cells
        
        if len(cells) < 4:
            continue
        
        # Читаем базовую информацию
        day_cell = cells[0].text.strip().lower()
        lesson_num = cells[1].text.strip()
        week_type = cells[2].text.strip()
        time_cell = cells[3].text.strip()
        
        # Определяем день недели
        if day_cell in day_of_week_map:
            current_day = day_of_week_map[day_cell]
        
        if not current_day or not lesson_num:
            continue
        
        # Парсим время
        time_start, time_end = parse_time(time_cell)
        if not time_start:
            continue
        
        # Парсим занятия для каждой группы
        for group_name, columns in groups_dict.items():
            for col_info in columns:
                col_idx = col_info['column_index']
                subgroup = col_info['subgroup']
                
                if col_idx >= len(cells):
                    continue
                
                lesson_text = cells[col_idx].text.strip()
                if not lesson_text:
                    continue
                
                lesson_data = parse_lesson_cell(lesson_text)
                if not lesson_data:
                    continue
                
                lessons.append({
                    'group_name': group_name,
                    'subgroup': subgroup,
                    'day_of_week': current_day,
                    'lesson_number': int(lesson_num) if lesson_num.isdigit() else None,
                    'time_start': time_start,
                    'time_end': time_end,
                    'week_type': week_type if week_type in ['Чет', 'Неч', 'Обе'] else 'Обе',
                    'subject': lesson_data['subject'],
                    'type': lesson_data['type'],
                    'teacher': lesson_data['teacher'],
                    'room': lesson_data['room'],
                    'start_date': lesson_data.get('start_date'),
                    'end_date': lesson_data.get('end_date')
                })
    
    return lessons

def main():
    file_path = 'schedules/23PG06_09.docx'
    
    if not Path(file_path).exists():
        print(f"❌ Файл {file_path} не найден")
        return
    
    print(f"📖 Парсинг файла: {file_path}")
    lessons = parse_schedule_docx(file_path)
    
    print(f"\n✅ Найдено занятий: {len(lessons)}")
    
    # Выводим примеры
    print("\n📚 Примеры занятий:")
    for lesson in lessons[:5]:
        print(f"\n  Группа: {lesson['group_name']} (подгруппа {lesson['subgroup']})")
        print(f"  День: {lesson['day_of_week']}, пара {lesson['lesson_number']}")
        print(f"  Время: {lesson['time_start']} - {lesson['time_end']}")
        print(f"  Неделя: {lesson['week_type']}")
        print(f"  Предмет: {lesson['subject']}")
        print(f"  Тип: {lesson['type']}")
        print(f"  Преподаватель: {lesson['teacher']}")
        print(f"  Аудитория: {lesson['room']}")

if __name__ == '__main__':
    main()
