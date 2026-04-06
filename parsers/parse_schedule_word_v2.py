#!/usr/bin/env python3
"""
Парсер расписания КГАСУ v2 - переписан с нуля
Правильно обрабатывает:
- Кабинеты и преподавателей
- Подгруппы
- Даты начала и окончания
- Чет/Неч недели
"""
import re
from pathlib import Path
from docx import Document
from datetime import datetime

def parse_time(time_str):
    """Парсит время из строки '8:00:00 - 9:30:00' или '8.00 - 9.30'"""
    if not time_str or '-' not in time_str:
        return None, None
    
    parts = time_str.split('-')
    if len(parts) != 2:
        return None, None
    
    start = parts[0].strip().replace('.', ':')
    end = parts[1].strip().replace('.', ':')
    
    # Добавляем :00 для секунд если их нет
    if start.count(':') == 1:
        start += ':00'
    if end.count(':') == 1:
        end += ':00'
    
    return start, end

def extract_dates(text):
    """
    Извлекает даты начала и окончания из текста
    Форматы:
    - "до 22.03." или "до 22. 03."
    - "С 06.04.26." или "с 06.04." или "С 06.04 по"
    """
    start_date = None
    end_date = None
    
    # Дата окончания: "до XX.XX.XX" или "до XX. XX."
    end_match = re.search(r'до\s+(\d{2})\s*\.\s*(\d{2})\s*\.?\s*(\d{2})?', text, re.IGNORECASE)
    if end_match:
        day = end_match.group(1)
        month = end_match.group(2)
        year = end_match.group(3) if end_match.group(3) else '26'
        end_date = f'20{year}-{month}-{day}'
    
    # Дата начала: "С XX.XX.XX." или "с XX.XX." или "С XX.XX по"
    start_patterns = [
        r'[Сс]\s+(\d{2})\.(\d{2})\.(\d{2})\.',  # С 06.04.26.
        r'[Сс]\s+(\d{2})\.(\d{2})\.',            # с 06.04.
        r'[Сс]\s+(\d{2})\.(\d{2})\s+по',         # С 06.04 по
    ]
    
    for pattern in start_patterns:
        start_match = re.search(pattern, text)
        if start_match:
            day = start_match.group(1)
            month = start_match.group(2)
            year = start_match.group(3) if len(start_match.groups()) >= 3 and start_match.group(3) else '26'
            start_date = f'20{year}-{month}-{day}'
            break
    
    return start_date, end_date

def parse_lesson_cell(cell_text):
    """
    Парсит ячейку с занятием
    
    Форматы:
    1. "Математика (Лекции) проф. Тимергалиев С. Н.2-203"
    2. "Математика (Лекции)доц. Гумеров А. В.3-208"
    3. "Математика (Лекции)доц. Гумеров А. В.  3-208"
    4. "Физкультура (Практические) СК «Тезуче» доц. Калманович В.Л."
    5. "Химия (Практические) до 12.04. доц. Валиева А. Ф. 1-36"
    6. "Химия (Практические) С 20.04. доц. Валиева А. Ф. 1-36"
    
    Возвращает: {subject, type, teacher, room, start_date, end_date}
    """
    if not cell_text or not cell_text.strip():
        return None
    
    text = cell_text.strip()
    
    # Извлекаем даты
    start_date, end_date = extract_dates(text)
    
    # Удаляем даты из текста для дальнейшего парсинга
    text = re.sub(r'до\s+\d{2}\s*\.\s*\d{2}\s*\.?\s*\d{0,2}\.?', '', text, flags=re.IGNORECASE)
    text = re.sub(r'[Сс]\s+\d{2}\.\d{2}\.?\d{0,2}\.?\s*(по)?', '', text)
    text = text.strip()
    
    # Извлекаем предмет и тип занятия: "Предмет (Тип)"
    match = re.match(r'(.+?)\s*\((.+?)\)', text)
    if not match:
        return {
            'subject': text,
            'type': None,
            'teacher': None,
            'room': None,
            'start_date': start_date,
            'end_date': end_date
        }
    
    subject = match.group(1).strip()
    lesson_type = match.group(2).strip()
    
    # Остальной текст после скобок
    rest = text[match.end():].strip()
    
    # Ищем аудиторию
    room = None
    
    # Специальные места (СК, НОЦ, в кавычках)
    special_room_match = re.search(r'(СК\s*[«"]?[^»"]+[»"]?|НОЦ\s*[«"]?[^»"]+[»"]?|[«"][^»"]+[»"])', rest)
    if special_room_match:
        room = special_room_match.group(1).strip()
        room = room.replace('«', '').replace('»', '').replace('"', '').strip()
        # Убираем из текста
        rest = rest[:special_room_match.start()].strip() + ' ' + rest[special_room_match.end():].strip()
        rest = rest.strip()
    
    # Обычная аудитория: цифры с корпусом (2-203) или без (317)
    # Ищем в конце строки или перед концом с пробелами
    if not room:
        # Паттерн: цифра-цифры или просто цифры, может быть с буквой
        room_patterns = [
            r'(\d+-\d+[А-Яа-яA-Za-z]?)\s*$',  # 2-203 или 2-203А в конце
            r'\s+(\d+-\d+[А-Яа-яA-Za-z]?)\s*',  # 2-203 с пробелами
            r'(\d{3,4}[А-Яа-яA-Za-z]?)\s*$',  # 317 или 317А в конце
            r'\s+(\d{3,4}[А-Яа-яA-Za-z]?)\s*',  # 317 с пробелами
        ]
        
        for pattern in room_patterns:
            room_match = re.search(pattern, rest)
            if room_match:
                room = room_match.group(1).strip()
                # Убираем аудиторию из текста
                rest = rest[:room_match.start()].strip() + rest[room_match.end():].strip()
                rest = rest.strip()
                break
    
    # Остальное - преподаватель
    teacher = rest.strip() if rest.strip() else None
    
    # Очистка преподавателя от лишних символов
    if teacher:
        # Убираем множественные пробелы
        teacher = re.sub(r'\s+', ' ', teacher)
        # Убираем точки в конце
        teacher = teacher.rstrip('.')
    
    return {
        'subject': subject,
        'type': lesson_type,
        'teacher': teacher,
        'room': room,
        'start_date': start_date,
        'end_date': end_date
    }

def parse_schedule_docx(file_path):
    """Парсит расписание из Word файла"""
    doc = Document(file_path)
    
    if len(doc.tables) == 0:
        print("❌ Таблицы не найдены")
        return []
    
    # Ищем таблицу с расписанием (должна быть достаточно большой)
    table = None
    for t in doc.tables:
        if len(t.rows) >= 3:
            table = t
            break
    
    if not table:
        print("❌ Таблица с расписанием не найдена")
        return []
    
    # Строка 0: заголовки с названиями групп
    # Строка 1: подгруппы
    # Строка 2+: данные
    
    header_row = table.rows[0]
    subgroup_row = table.rows[1]
    
    # Собираем информацию о группах и колонках
    groups_info = {}  # {group_name: [(col_idx, subgroup), ...]}
    
    for col_idx, cell in enumerate(header_row.cells):
        if col_idx < 4:  # Пропускаем служебные колонки
            continue
        
        group_name = cell.text.strip()
        
        # Проверяем что это название группы (формат: 25СЖ01)
        if not group_name or not re.match(r'\d{1,2}[А-ЯЁ]{2}\d{2}', group_name):
            continue
        
        # Определяем подгруппу
        subgroup_text = subgroup_row.cells[col_idx].text.strip().lower()
        subgroup = 0  # По умолчанию для всей группы
        
        if 'подгруппа 1' in subgroup_text or 'подгр. 1' in subgroup_text or 'подгр.1' in subgroup_text:
            subgroup = 1
        elif 'подгруппа 2' in subgroup_text or 'подгр. 2' in subgroup_text or 'подгр.2' in subgroup_text:
            subgroup = 2
        elif 'подгруппа 3' in subgroup_text:
            subgroup = 3
        elif 'подгруппа 4' in subgroup_text:
            subgroup = 4
        
        if group_name not in groups_info:
            groups_info[group_name] = []
        
        # Проверяем что не добавляем дубликат (группа + подгруппа уже есть)
        if not any(sg == subgroup for _, sg in groups_info[group_name]):
            groups_info[group_name].append((col_idx, subgroup))
    
    print(f"📋 Найдено групп: {len(groups_info)}")
    for group_name, cols in groups_info.items():
        subgroups_str = ', '.join([f"подгр.{sg}" if sg > 0 else "вся группа" for _, sg in cols])
        print(f"  - {group_name}: {subgroups_str}")
    
    # Парсим расписание
    lessons = []
    
    day_map = {
        'понедельник': 1,
        'вторник': 2,
        'среда': 3,
        'четверг': 4,
        'пятница': 5,
        'суббота': 6,
        'воскресенье': 7
    }
    
    current_day = None
    
    for row_idx in range(2, len(table.rows)):
        row = table.rows[row_idx]
        cells = row.cells
        
        if len(cells) < 4:
            continue
        
        # Читаем служебные колонки
        day_cell = cells[0].text.strip().lower()
        lesson_num_cell = cells[1].text.strip()
        week_type_cell = cells[2].text.strip()
        time_cell = cells[3].text.strip()
        
        # Определяем день недели
        if day_cell in day_map:
            current_day = day_map[day_cell]
        
        if not current_day:
            continue
        
        # Проверяем номер пары
        if not lesson_num_cell or not lesson_num_cell.isdigit():
            continue
        
        lesson_number = int(lesson_num_cell)
        
        # Парсим время
        time_start, time_end = parse_time(time_cell)
        if not time_start:
            continue
        
        # Определяем тип недели
        week_type = 'Обе'
        if week_type_cell in ['Чет', 'Неч']:
            week_type = week_type_cell
        
        # Парсим занятия для каждой группы
        for group_name, columns in groups_info.items():
            for col_idx, subgroup in columns:
                if col_idx >= len(cells):
                    continue
                
                lesson_text = cells[col_idx].text.strip()
                
                if not lesson_text:
                    continue
                
                lesson_data = parse_lesson_cell(lesson_text)
                
                if not lesson_data or not lesson_data['subject']:
                    continue
                
                lessons.append({
                    'group_name': group_name,
                    'subgroup': subgroup,
                    'day_of_week': current_day,
                    'lesson_number': lesson_number,
                    'time_start': time_start,
                    'time_end': time_end,
                    'week_type': week_type,
                    'subject': lesson_data['subject'],
                    'type': lesson_data['type'],
                    'teacher': lesson_data['teacher'],
                    'room': lesson_data['room'],
                    'start_date': lesson_data.get('start_date'),
                    'end_date': lesson_data.get('end_date')
                })
    
    return lessons

def main():
    """Тестирование парсера"""
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = '../schedules/25SZH01_02.docx'
    
    if not Path(file_path).exists():
        print(f"❌ Файл {file_path} не найден")
        return
    
    print(f"📖 Парсинг файла: {file_path}")
    lessons = parse_schedule_docx(file_path)
    
    print(f"\n✅ Найдено занятий: {len(lessons)}")
    
    # Группируем по группам
    by_group = {}
    for lesson in lessons:
        group = lesson['group_name']
        if group not in by_group:
            by_group[group] = []
        by_group[group].append(lesson)
    
    print(f"\n📊 По группам:")
    for group, group_lessons in by_group.items():
        print(f"  {group}: {len(group_lessons)} занятий")
    
    # Показываем примеры
    print(f"\n📚 Примеры занятий:")
    for lesson in lessons[:5]:
        print(f"\n  Группа: {lesson['group_name']} (подгр. {lesson['subgroup']})")
        print(f"  День: {lesson['day_of_week']}, пара {lesson['lesson_number']}, {lesson['week_type']}")
        print(f"  Время: {lesson['time_start']} - {lesson['time_end']}")
        print(f"  Предмет: {lesson['subject']}")
        print(f"  Тип: {lesson['type']}")
        print(f"  Преподаватель: {lesson['teacher']}")
        print(f"  Аудитория: {lesson['room']}")
        if lesson['start_date']:
            print(f"  Начало: {lesson['start_date']}")
        if lesson['end_date']:
            print(f"  Конец: {lesson['end_date']}")

if __name__ == '__main__':
    main()
